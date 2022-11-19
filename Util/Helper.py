import os.path
from datetime import datetime, date
from openpyxl import load_workbook
from docx import Document

from Util.Person import Person


def getValuesFromExcel():
    path = os.path.abspath(os.curdir)
    wb = load_workbook(path+"\\Infoinput.xlsx")
    ws = wb.active

    kasualie = (ws["B3"].value)

    taufling = Person(ws["B7"].value, ws["C7"].value, "evangelisch")
    if type(ws["B8"].value) != str:                 #only if no string ...
        taufling.dateOfBirth = datetime.strftime((ws["B8"].value), "%d/%b/%Y")
    taufling.placeOfBirth = ws["B9"].value


    eltern1 = Person(ws["B12"].value, ws["C12"].value, ws["B14"].value)
    eltern1.originName = " " if ws["B13"].value == "EGNAME01" else ", geb. "+(ws["B13"].value)
    eltern2 = Person(ws["B16"].value, ws["C16"].value, ws["B18"].value)
    eltern2.originName = " " if ws["B17"].value == "EGNAME02" else ", geb. "+(ws["B17"].value)

    global listOfTaufpate
    listOfTaufpate= []
    global listOfTaufzeuge
    listOfTaufzeuge= []


    # check wether its pate or zeuge
    def addToPateOrZeugeList(pateOderZeuge, confession):

        if (confession == "ohne"):
           listOfTaufzeuge.append(pateOderZeuge)
        else:
            listOfTaufpate.append(pateOderZeuge)

    dummyPerson = Person("", "", "")
    if (ws["B21"].value != "PNNAME01"):
        pateOrZeuge1 = Person(ws["B21"].value, ws["C21"].value, ws["B22"].value)
    else:
        pateOrZeuge1 = dummyPerson
    if (ws["B24"].value) != "PNNAME02":
        pateOrZeuge2 = Person(ws["B24"].value, ws["C24"].value, ws["B25"].value)
    else:
        pateOrZeuge2 = dummyPerson
    if (ws["B27"].value) != "PNNAME03":
        pateOrZeuge3 = Person(ws["B27"].value, ws["C27"].value, ws["B28"].value)
    else:
        pateOrZeuge3 = dummyPerson
    if (ws["B30"].value) != "PNNAME04":
        pateOrZeuge4 = Person(ws["B30"].value, ws["C30"].value, ws["B31"].value)
    else:
        pateOrZeuge4 = dummyPerson

    addToPateOrZeugeList(pateOrZeuge1, pateOrZeuge1.confession)
    addToPateOrZeugeList(pateOrZeuge2, pateOrZeuge2.confession)
    addToPateOrZeugeList(pateOrZeuge3, pateOrZeuge3.confession)
    addToPateOrZeugeList(pateOrZeuge4, pateOrZeuge4.confession)

    kdatum = date.today()
    if type(ws["B34"].value) != str:
        kdatum = datetime.strftime((ws["B34"].value), "%d/%b/%Y")

    bspruch = (ws["B35"].value)
    bstelle = (ws["B36"].value)
    pfarrer = (ws["B37"].value)

    adatum = date.today()
    if type(ws["B39"].value) != str:
        adatum = datetime.strftime((ws["B39"].value), "%d/%b/%Y")

    aort = (ws["B40"].value)

    allValuesDict = {
        "kasualie": kasualie,
        "kdatum": kdatum,
        "täufling": taufling,
        "eltern": [eltern1, eltern2],
        "taufpate": listOfTaufpate,
        "taufzeuge": listOfTaufzeuge,
        "spruch": bspruch,
        "stelle": bstelle,
        "pfarrer": pfarrer,
        "adatum": adatum,
        "aort": aort
    }

    return allValuesDict


def checkKasualie(myDict):
    kasualie = myDict.get("kasualie")
    if kasualie == "Trauung":
        for filename in os.listdir("Formulare"):
            if "Trauung" in filename:
                fillFormulare(filename, myDict)
    elif kasualie == "Taufe":
        for filename in os.listdir("Formulare"):
            if "Trauung" not in filename:
                fillFormulare(filename, myDict)



def fillFormulare(file, dataDict):
    formdoc = Document("."+"\\Formulare\\" + file)
    täufling = dataDict.get("täufling")
    if "StammbuchTaufe" in file:
        fillConcreteFormular(formdoc, dataDict, 0)
        formdoc.save('Stammbuch_Taufe ' + täufling.lastname + '.docx')
    elif "StammbuchTrauung" in file:
        ename = dataDict.get("eltern")
        fillConcreteFormular(formdoc, dataDict, 0)
        formdoc.save('Stammbuch_Trauung ' + ename[0].lastname + '.docx')
    if "Patenurkunde" in file:
        pate = dataDict.get("taufpate")

        for idx in range(len(pate)):
            formdoc = Document("."+"\\Formulare\\" + file)
            fillConcreteFormular(formdoc, dataDict, idx)
            formdoc.save('Patenurkunde ' + täufling.lastname + "_" + pate[idx].lastname + '.docx')

    if "Taufzeuge" in file:
        zeuge = dataDict.get("taufzeuge")
        zeugeIdx = 0
        for concreteZeuge in zeuge:
            formdoc = Document("."+"\\Formulare\\"+ file)
            fillConcreteFormular(formdoc, dataDict, zeugeIdx)
            formdoc.save('Taufzeuge ' + täufling.lastname + "_" + concreteZeuge.lastname + '.docx')
            zeugeIdx += 1
    if "Taufurkunde" in file:
        fillConcreteFormular(formdoc, dataDict, 0)
        formdoc.save('Taufurkunde ' + täufling.lastname + '_' + täufling.firstname +'.docx')


def fillConcreteFormular(formdoc, dataDict, idx):
    # TODO ersetzen durch Template Names
    lplaceh = ["TVNAME", "TNNAME", "GDATUM", "GORT",
               "EVNAME01", "ENNAME01", "EGNAME01", "E01BEKENNTNIS",
               "EVNAME02", "ENNAME02", "EGNAME02", "E02BEKENNTNIS",
               "PVNAME01", "PNNAME01",
               "PVNAME02", "PNNAME02",
               "PVNAME03", "PNNAME03",
               "ZVNAME01", "ZNNAME01",
               "KDATUM", "BSPRUCH", "BSTELLE", "PFARRER", "ADATUM",
               ]
    taufling = dataDict.get("täufling")
    eltern = dataDict.get("eltern")
    paten = dataDict.get("taufpate")
    zeuge = dataDict.get("taufzeuge")

    pate2fn = paten[1].firstname if len(paten) > 1 else ""
    pate2ln = paten[1].lastname if len(paten) > 1 else ""
    pate3fn = paten[2].firstname if len(paten) > 2 else ""
    pate3ln = paten[2].lastname if len(paten) > 2 else ""


    lreplace = [taufling.firstname, taufling.lastname, taufling.dateOfBirth, taufling.placeOfBirth,
                eltern[0].firstname, eltern[0].lastname, eltern[0].originName, eltern[0].confession,
                eltern[1].firstname, eltern[1].lastname, eltern[1].originName, eltern[1].confession,
                paten[idx].firstname, paten[idx].lastname,
                pate2fn, pate2ln,
                pate3fn, pate3ln,
                zeuge[idx].firstname, zeuge[idx].lastname,
                dataDict.get("kdatum"), dataDict.get("spruch"), dataDict.get("stelle"), dataDict.get("pfarrer"), dataDict.get("adatum")
                ]
    for paragraph in formdoc.paragraphs:
        for idx, var in enumerate(lplaceh):
            if var in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if var in inline[i].text:
                        text = inline[i].text.replace(str(var), lreplace[idx])
                        inline[i].text = text



