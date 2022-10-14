print ("Marta\n\nLukas 10,41f. 'Marta, Marta, du hast viel Sorge und Mühe. Eins aber ist not. Maria hat das gute Teil erwählt.'\n\n")



import os
import datetime
import random

import openpyxl                 # to work with excel
from openpyxl import load_workbook
from docx import Document       # to work with Document
from pathlib import Path


path = "."
document = Document()
formdoc = Document()

# Vars to check ...
#Täufling
tnname = "TNNAME"
tvname = "TVNAME"
gdatum = "GDATUM"
gort = "GORT"

#Eltern/Brautpaar
enname01 = "ENNAME01"
evname01 = "EVNAME01"
egname01 = "EGNAME01"
e01bekenntnis = "E01BEKENNTNIS"

enname02 = "ENNAME02"
evname02 = "EVNAME02"
egname02 = "EGNAME02"
e02bekenntnis = "E02BEKENNTNIS"

#Taufpaten/Zeugen
pnname01 = "PNNAME01"
pvname01 = "PVNAME01"
p01bekenntnis = "P01BEKENNTNIS"

pnname02 = "PNNAME02"
pvname02 = "PVNAME02"
p02bekenntnis = "P02BEKENNTNIS"

pnname03 = "PNNAME03"
pvname03 = "PVNAME03"
p03bekenntnis = "P03BEKENNTNIS"

pnname04 = "PNNAME04"
pvname04 = "PVNAME04"
p04bekenntnis = "P04BEKENNTNIS"

#zur Kasualie
kdatum = "KDATUM"
bspruch = "BSPRUCH"
bstelle = "BSTELLE"
pfarrer = "PFARRER"

adatum = "ADATUM"
aort = "AORT"

# Erstelle eine placeholder Liste, von allen placeholdern, die in den Formularen vorhanden sind um später lplaceh durch lvars zu ersetzen

lplaceh = [tnname, tvname, gdatum, gort, enname01, evname01, egname01, e01bekenntnis, enname02, evname02, egname02, e02bekenntnis, pnname01, pvname01, p01bekenntnis, pnname02, pvname02, p02bekenntnis, pnname03, pvname03, p03bekenntnis, pnname04, pvname04, p04bekenntnis, kdatum, bspruch, bstelle, pfarrer, adatum, aort]


def getTheVars():
    print("This are all the Infos we got:")
    global path
    wb = load_workbook(path+"\\Infoinput.xlsx")
    #name= #!/usr/bin/python
    ws = wb.active

    global kasualie
    kasualie = (ws["B3"].value)

    global tnname
    tnname = (ws["B7"].value)
    global tvname
    tvname = (ws["C7"].value)
    global gdatum
    if type(ws["B8"].value) != str:                 #only if no string ...
        gdatum = datetime.datetime.strftime((ws["B8"].value), "%d/%b/%Y")
    global gort
    gort = (ws["B9"].value)
    global enname01
    enname01 = (ws["B12"].value)
    global evname01
    evname01 = (ws["C12"].value)
    global egname01
    if (ws["B13"].value) == "EGNAME01":
        egname01 = " "
    else:
        egname01 = ", geb. "+(ws["B13"].value)
    global e01bekenntnis
    e01bekenntnis = (ws["B14"].value)

    global enname02
    enname02 = (ws["B16"].value)
    global evname02
    evname02 = (ws["C16"].value)
    global egname02
    if (ws["B17"].value) == "EGNAME02":
        egname02 = " "
    else:
        egname02 = ", geb. "+(ws["B17"].value)
    global e02bekenntnis
    e02bekenntnis = (ws["B18"].value)

    global tplist, tzlist       # creating a list for "Taufpaten" and "Taufzeugen"
    #tplist = [[],[]]
    #tzlist = [[],[]]
    tplist = [[] for x in range(4)]
    tzlist = [[] for x in range(4)]
#    tzlist = []

    global pnname01, pvname01, p01bekenntnis
    if (ws["B21"].value) == "PNNAME01":
        pnname01 = " "
        pvname01 = " "
    else:
        pnname01 = (ws["B21"].value)
        pvname01 = (ws["C21"].value)
        p01bekenntnis = (ws["B22"].value)
        if p01bekenntnis == "ohne":
    #        tzlist.append(pnname01)
            tzlist[0].append(pnname01)
            tzlist[0].append(pvname01)
            tzlist[0].append(p01bekenntnis)
        else:
            tplist[0].append(pnname01)
            tplist[0].append(pvname01)
            tplist[0].append(p01bekenntnis)
    print(tzlist)
    print("adawdeawdawdawd")
#    print(tplist)


    global pnname02, pvname02, p02bekenntnis
    if (ws["B24"].value) == "PNNAME02":
        pnname02 = " "
        pvname02 = " "
    else:
        pnname02 = (ws["B24"].value)
        pvname02 = (ws["C24"].value)
        p02bekenntnis = (ws["B25"].value)
        if p02bekenntnis == "ohne":
        #    tzlist.append(pnname02)
            tzlist[1].append(pnname02)
            tzlist[1].append(pvname02)
            tzlist[1].append(p02bekenntnis)
        else:
            tplist[1].append(pnname02)
            tplist[1].append(pvname02)
            tplist[1].append(p02bekenntnis)

    global pnname03, pvname03, p03bekenntnis
    if (ws["B27"].value) == "PNNAME03":
        pnname03 = " "
        pvname03 = " "
    else:
        pnname03 = (ws["B27"].value)
        pvname03 = (ws["C27"].value)
        p03bekenntnis = (ws["B28"].value)
        if p03bekenntnis == "ohne":
        #    tzlist.append(pnname03)
            tzlist[2].append(pnname03)
            tzlist[2].append(pvname03)
            tzlist[2].append(p03bekenntnis)
        else:
            tplist[2].append(pnname03)
            tplist[2].append(pvname03)
            tplist[2].append(p03bekenntnis)

    global pnname04, pvname04, p04bekenntnis
    if (ws["B30"].value) == "PNNAME04":
        pnname04 = " "
        pvname04 = " "
    else:
        pnname04 = (ws["B30"].value)
        pvname04 = (ws["C30"].value)
        p04bekenntnis = (ws["B31"].value)
        if p04bekenntnis == "ohne":
            #tzlist.append(pnname04)
            tzlist[3].append(pnname04)
            tzlist[3].append(pvname04)
            tzlist[3].append(p04bekenntnis)
        else:
            tplist[3].append(pnname04)
            tplist[3].append(pvname04)
            tplist[3].append(p04bekenntnis)


    global kdatum
    if type(ws["B34"].value) != str:
        kdatum = datetime.datetime.strftime((ws["B34"].value), "%d/%b/%Y")
    global bspruch
    bspruch = (ws["B35"].value)
    global bstelle
    bstelle = (ws["B36"].value)
    global pfarrer
    pfarrer = (ws["B37"].value)

    global adatum
    if type(ws["B39"].value) != str:
        adatum = datetime.datetime.strftime((ws["B39"].value), "%d/%b/%Y")
    global aorts
    aort = (ws["B40"].value)


    global lvars                # creating a list of all the vars
    lvars = [tnname, tvname, gdatum, gort, enname01, evname01, egname01, e01bekenntnis, enname02, evname02, egname02, e02bekenntnis, pnname01, pvname01, p01bekenntnis, pnname02, pvname02, p02bekenntnis, pnname03, pvname03, p03bekenntnis, pnname04, pvname04, p04bekenntnis, kdatum, bspruch, bstelle, pfarrer, adatum, aort]            #put new vars here ... also put them in the lplaceh list ....

    print(tplist)
    print(tzlist)





#    global clist
#    clist = ["B3","C3","B4","B5","B6","B8","B9","B10","B13","B14","B15","B16","B19","B20","B21","B23","B24"]          #hier kommen alle Excell felder rein, die am ende gecleart werden sollen. ...
#    tabclear(ws)
#    wb.save("Infoinput.xlsx")

#def tabclear(tab):               #ich sollte die Excell Tabelle clearen ...
#    for i in clist:
#        tab[i].value = ''

def paraMove(output_doc_name, paragraph):           # to keep the style
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        output_run.bold = run.bold                                      # Run's bold data
        output_run.italic = run.italic                                  # Run's italic data
        output_run.underline = run.underline                            # Run's underline data
        output_run.font.color.rgb = run.font.color.rgb                  # Run's color data
        output_run.style.name = run.style.name                          # Run's font data
        output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment


#def brickMove(doc):
#    #print("Brick Moved")
#    global document
#    input_doc = Document(path+doc)
#    for para in input_doc.paragraphs:
#        paraMove(document, para)


#def fillVars():
#    print("\nFilling the Vars")
#    lpos=-1
#
#    for var in lplaceh:
#        lpos += 1                   #check list positions
#        for paragraph in document.paragraphs:
#            if var in paragraph.text:
#                paragraph.text = paragraph.text.replace(str(var),lvars[lpos])

def fillFormular(file):                             # Das Formular wird geöffnet und mit den entsprechenden Variablen befüllt.
    print("\nFilling te vars into: "+file+" ...")
    formdoc = Document("."+"\\Formulare\\"+ file)

    # todo: für jeden Paten ein Dokument
    # for  tplist.len() > 1:   für jeden paten und jeden zeugen for pate in patenliste ... fülle eine Patenurkunde aus und speicher das ..


    #Speichern entsprechend der Urkunde/des Formulars
#    if "Patenurkunde" in file:
#        formdoc.save('Patenurkunde '+tvname +'_'+tnname+'.docx')
    if "StammbuchTaufe" in file:
        filling(formdoc)
        print("\n\nFormular wurde erstellt ...")
        formdoc.save('Stammbuch_Taufe '+tvname +'_'+tnname+'.docx')
    elif "StammbuchTrauung" in file:
        print("testtest")
        filling(formdoc)
        print("\n\nFormular wurde erstellt ...")
        formdoc.save('Stammbuch_Trauung '+enname01+'.docx')
    elif "Taufurkunde" in file:
        filling(formdoc)
        print("\n\nFormular wurde erstellt ...")
        formdoc.save('Taufurkunde '+tvname +'_'+tnname+'.docx')
#    elif "Taufzeuge" in file:
#        formdoc.save('Taufzeuge '+tvname +'_'+tnname+'.docx')

    for pate in tplist:
        if "Patenurkunde" in file:
            formdoc = Document("."+"\\Formulare\\"+ file)
            lvars[12] =  pate[0]
            lvars[13] =  pate[1]
            filling(formdoc)
            print("\n\nFormular wurde erstellt ...")
            formdoc.save('Patenurkunde '+tnname +'_'+pate[0]+'.docx')
    for zeuge in tzlist:
        if "Taufzeuge" in file:
            formdoc = Document("."+"\\Formulare\\"+ file)
            lvars[12] =  zeuge[0]
            lvars[13] =  zeuge[1]
            filling(formdoc)
            print("\n\nFormular wurde erstellt ...")
            formdoc.save('Taufzeuge '+tnname +'_'+zeuge[0]+'.docx')

def filling(formdoc):
    lpos=0
    for var in lplaceh:
        for paragraph in formdoc.paragraphs:
            if var in paragraph.text:
                print(var)
    #            paragraph.text = paragraph.text.replace(str(var),lvars[lpos])
                inline = paragraph.runs                     # Das braucht man wohl, damit die Formatierung erhalten bleibt, durchblicke ich nicht ganz
                for i in range(len(inline)):
                    if var in inline[i].text:
                        print(str(var))

                    #    print(inline[i].text)
                        print("wichtige infos yeah")
                        print(lvars[lpos])
                        text = inline[i].text.replace(str(var),lvars[lpos])
                        inline[i].text = text
        lpos += 1  #inc list position

#lplaceh = [tnname, tvname, gdatum, gort, enname01, evname01, egname01, e01bekenntnis, enname02, evname02, egname02, e02bekenntnis, pnname01, pvname01, p01bekenntnis, pnname02, pvname02, p02bekenntnis, pnname03, pvname03, p03bekenntnis, pnname04, pvname04, p04bekenntnis, kdatum, bspruch, bstelle, pfarrer, adatum, aort]



try:
    getTheVars()                        # get vars from the Infoinput
    #fillVars()                          # filling the doc with the right vars (placeholder into vars)
    #document.save('Beerdigung '+lname+'.docx')        # just save the final doc
    #print("\n\nDokument wurde erstellt")
    if kasualie == "Trauung":
        print("Aha eine Trauung ...")
        for filename in os.listdir("Formulare"):
            if "Trauung" in filename:
                fillFormular(filename)
    elif kasualie == "Taufe":
        print("Aha eine Taufe ...")
        for filename in os.listdir("Formulare"):
            if "Trauung" not in filename:
                fillFormular(filename)
    #    with open(os.path.join("files", filename), 'r') as f:
    #        text = f.read()
    #        print(text)

finally:
    input("\nBis zum nächsten Mal ....")




# ------ ToDo:  --------------------------------------------------------------------------------------------
#    -erledigt-     mehrere Paten möglich
#    -erledigt-     nur noch relevante Urkunden errstellen
#                   Paten und Taufzeugen unterscheiden ... aber erstmal weiter ...
#                   Stammbuch schauen, dass die formatierung nicht zerschossen wird
#                   Trauurkunde fehlt ...
#                   Konfirmation ergänzen ... muss auch unten bei der "Taufe" eingefügt werden ...
#                   tplist und tzlist als erweiterpare Liste definieren ...
#
#
#
#
#
#
#
#
#
#
